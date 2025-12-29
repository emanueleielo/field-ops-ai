"""DOCX text extraction using python-docx."""

from io import BytesIO
from pathlib import Path

from docx import Document

from app.services.extractors.base import (
    BaseExtractor,
    ExtractionError,
    ExtractionResult,
)


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX documents."""

    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            ExtractionResult containing extracted text.

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            doc = Document(file_path)
            return self._extract_from_document(doc)
        except Exception as e:
            raise ExtractionError(f"Failed to extract text from DOCX: {e}") from e

    def extract_from_bytes(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract text from DOCX bytes.

        Args:
            content: DOCX content as bytes.
            filename: Original filename (unused for DOCX).

        Returns:
            ExtractionResult containing extracted text.

        Raises:
            ExtractionError: If extraction fails.
        """
        try:
            doc = Document(BytesIO(content))
            return self._extract_from_document(doc)
        except Exception as e:
            raise ExtractionError(f"Failed to extract text from DOCX: {e}") from e

    def _extract_from_document(self, doc: Document) -> ExtractionResult:
        """Extract text from a Document instance.

        Args:
            doc: python-docx Document instance.

        Returns:
            ExtractionResult containing extracted text.
        """
        paragraphs: list[str] = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if it's a heading
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
            if table_rows:
                paragraphs.append("\n[Table]\n" + "\n".join(table_rows) + "\n")

        full_text = "\n".join(paragraphs)

        # Extract metadata from core properties
        metadata: dict[str, str] = {}
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = str(doc.core_properties.title)
            if doc.core_properties.author:
                metadata["author"] = str(doc.core_properties.author)
            if doc.core_properties.subject:
                metadata["subject"] = str(doc.core_properties.subject)

        return ExtractionResult(
            text=full_text,
            page_count=1,  # DOCX doesn't have a direct page count
            metadata=metadata,
        )
